# distance_exporter.py
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DistanceExporter:
    """Класс для экспорта таблицы расстояний в формат Word"""

    @staticmethod
    def export_to_word(distances: dict, objects: list, filename: str):
        """
        Экспортирует таблицу расстояний в документ Word

        Args:
            distances: словарь расстояний между объектами
            objects: список объектов
            filename: путь для сохранения файла
        """
        doc = Document()

        # Добавляем заголовок
        heading = doc.add_paragraph("Таблица расстояний между объектами")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Создаем таблицу
        n = len(objects)
        table = doc.add_table(rows=n + 1, cols=n + 1)
        table.style = 'Table Grid'

        # Заполняем заголовки
        table.cell(0, 0).text = "Наименование"
        for i, obj in enumerate(objects):
            table.cell(0, i + 1).text = obj.name
            table.cell(i + 1, 0).text = obj.name

        # Заполняем значения
        for i, obj1 in enumerate(objects):
            for j, obj2 in enumerate(objects):
                if obj1.id == obj2.id:
                    value = "-"
                else:
                    value = str(distances[obj1.id][obj2.id])
                table.cell(i + 1, j + 1).text = value

        # Форматируем таблицу
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Устанавливаем ширину столбцов
        for cell in table.columns[0].cells:
            cell.width = Inches(2.0)

        # Сохраняем документ
        doc.save(filename)